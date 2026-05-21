import os
from PIL import Image
import pytesseract

from config.settings import TESSERACT_PATH

pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH


def _extract_text_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    texts = []
    for page in doc:
        # Try direct text extraction first (fast, no OCR needed)
        text = page.get_text().strip()
        if text:
            texts.append(text)
        else:
            # Fallback to OCR for scanned/image-based pages
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            texts.append(pytesseract.image_to_string(img))
    doc.close()
    return "\n".join(texts)


def _extract_text_docx(file_path: str) -> str:
    # DOCX -> direct text extraction
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "DOCX upload requires python-docx. Run: pip install python-docx"
        ) from e

    doc = docx.Document(file_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    # Fallback: include table text too (best-effort)
    if not parts:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = (cell.text or "").strip()
                    if txt:
                        parts.append(txt)

    return "\n".join(parts)


def _extract_text_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_text_image(file_path: str) -> str:
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)


def extract_text(file_path: str) -> str:
    """Extract text from supported uploads.

    Supported:
      - .pdf  : PDF -> images -> OCR
      - .docx : direct extraction (no OCR)
      - .txt  : direct extraction
      - other : treated as image -> OCR
    """
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            return _extract_text_pdf(file_path)

        if ext == ".docx":
            return _extract_text_docx(file_path)

        if ext == ".txt":
            return _extract_text_txt(file_path)

        # Fallback: treat as image
        return _extract_text_image(file_path)

    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"Text extraction failed for {ext or 'file'}: {str(e)}")

